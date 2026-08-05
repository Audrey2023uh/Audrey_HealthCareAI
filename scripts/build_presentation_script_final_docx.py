#!/usr/bin/env python3
"""
Build a full graduate-level / conference-style Milestone 2 presentation transcript.
Not bullet speaker notes — continuous technical presentation prose for ~15–20 minutes.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(
    r"C:\Users\audre\OneDrive\1- Final project\AI HEALTH CARE\Milestone2_Presentation_Script_Final.docx"
)

NAVY = RGBColor(0x0F, 0x4C, 0x81)


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        font(r, size=16, bold=True, color=NAVY)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        font(r, size=13, bold=True, color=NAVY)


def para(doc, text, *, size=11, bold=False, italic=False, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    font(r, size=size, bold=bold, italic=italic)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # ========== COVER ==========
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Milestone 2 — Complete Presentation Transcript")
    font(r, size=20, bold=True, color=NAVY)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Source-Linked Clinical Evidence CDSS")
    font(r2, size=15, bold=True)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "A graduate-level AI engineering presentation script\n"
        "for approximately 15–20 minutes of continuous delivery\n\n"
        "Audrey Rah · University of Houston · Houston, Texas, USA\n"
        "Codebase folder: AI HEALTH CARE/\n"
        "Aligned with: proposal · Milestone 1 narrative · AI Tech Stack.docx · Chip Huyen–style AI engineering honesty"
    )
    font(r3, size=11)

    para(
        doc,
        "How to use this document: This is a full talk transcript, not slide captions. "
        "Speak it in your own voice. Pause at figure explanations. During the demo section, "
        "switch to the live Streamlit app while continuing the narrative. "
        "Throughout the talk, keep one discipline: yellow = implemented today; "
        "green = ideal production architecture that is not claimed as finished.",
        italic=True,
        after=14,
    )

    # ========== OPENING ==========
    h1(doc, "Opening — Setting the problem and the Milestone 2 contract")
    para(
        doc,
        "Good afternoon. My name is Audrey Rah, and I’m presenting Milestone 2 of my Healthcare AI project "
        "at the University of Houston: a source-linked Clinical Evidence Clinical Decision Support System.",
    )
    para(
        doc,
        "I want to begin with the clinical problem in plain language, then move quickly into engineering, "
        "because Milestone 2 is judged as an AI engineering deliverable — not as a product pitch.",
    )
    para(
        doc,
        "In outpatient and inpatient care, clinicians are under time pressure and face an enormous literature: "
        "clinical practice guidelines, systematic reviews, trials, government evidence resources, and patient-education material. "
        "General-purpose language models can answer medical questions fluently, but fluency is not the same as verifiability. "
        "In healthcare, an answer without a checkable source is a liability. Our project addresses that gap: "
        "help a clinician — for example a resident like Dr. Lin from our Milestone 1 persona — retrieve stronger evidence faster, "
        "show where every claim came from, and remain explicit when evidence is weak, conflicting, or when patient data is incomplete.",
    )
    para(
        doc,
        "Milestone 1 established the domain, the persona, the evidence hierarchy philosophy, and the principle that the clinician stays in charge. "
        "Milestone 2 is the engineering proof. I will show you a runnable system in the AI HEALTH CARE codebase: "
        "a Streamlit clinician workspace, a LangGraph multi-agent CDSS pipeline, FAISS-based retrieval-augmented generation, "
        "priority ranking from guidelines down to evaluation-only MedQuAD material, citation verification, confidence scoring, "
        "and an optional Knowledge Update path. Before any architecture detail, one sentence that must stay visible in every demo: "
        "this is a research and education prototype. It is not a medical device. A licensed clinician makes the final decision.",
    )
    para(
        doc,
        "My talk has five movements. First, what was actually built. Second, architecture figures — end-to-end layers, LangGraph edges, "
        "RAG, ranking, and module interactions. Third, the instructor AI Tech Stack table and the yellow-versus-green honesty contract. "
        "Fourth, trade-offs, evaluation, and lessons. Fifth, a live demo that proves the claims.",
    )

    # ========== SLIDE 3 CONTENT EXPANDED ==========
    h1(doc, "Part I — What was built, and what that means in AI engineering terms")
    h2(doc, "The honesty boundary")
    para(
        doc,
        "In Chip Huyen’s AI Engineering framing, a serious system presentation separates the application you can run today "
        "from the production architecture you would fund later. I treat that separation as a first-class deliverable. "
        "If I blur them, I am no longer doing engineering — I am doing aspiration theater. So when I say “implemented,” "
        "I mean there is executable Python in AI HEALTH CARE that you can run with streamlit run app.py after building the index. "
        "When I say “ideal,” I mean a production-oriented upgrade path that I am not claiming ships in this semester prototype.",
    )
    para(
        doc,
        "What was built, in resume language and in technical language at the same time, is three things.",
    )
    para(
        doc,
        "First, an end-to-end clinical decision-support prototype. The clinician interacts with a Streamlit UI: patient assessment fields, "
        "a clinical question, and a run of the CDSS workflow. Behind that UI is a LangGraph StateGraph with eight nodes. "
        "The pipeline returns a structured result: patient assessment, risk analysis, retrieved evidence hits, verification notes, "
        "recommendations with evidence summary, confidence, needs_human_review, timings, and an agent_trace. "
        "That end-to-end path is the Milestone 2 definition of “working system.”",
    )
    para(
        doc,
        "Second, hierarchy-aware RAG over a public-style medical knowledge base. Documents are curated into seed_medical_kb.json, "
        "cleaned and chunked, embedded with local TF-IDF by default, and stored in FAISS IndexFlatIP with chunk metadata. "
        "At query time we retrieve candidates, then re-rank hard by Priority 1 through 7. Semantic similarity is a recall signal, "
        "not the final clinical authority. Superseded guidelines are penalized. Generation — if an OpenAI-compatible key is present — "
        "is constrained to retrieved context with [n] citations. If no key is present, the system still runs using extractive grounded summaries. "
        "That offline fallback is not a compromise I hide; it is an engineering reliability feature for classroom demos.",
    )
    para(
        doc,
        "Third, explainability and measurement. The UI surfaces evidence category tiles, conflict flags, citations, confidence, and agent traces. "
        "scripts/evaluate.py writes retrieval and citation metrics to outputs/eval_results.json. FastAPI /ask is optional for API-shaped demos. "
        "A Knowledge Update Agent can probe authoritative URLs and apply local supersession rules, then rebuild the index. "
        "That is the implemented freshness loop — not a full automatic PDF ingest of every society release.",
    )

    h1(doc, "Part II — AI engineering layers: how to think about the system")
    para(
        doc,
        "Before the figures, I want you to hold a four-layer mental model, because every later diagram is just a zoom into one of these layers.",
    )
    para(
        doc,
        "Application layer: Streamlit clinician workspace, forms, query composition, result panels, disclaimer, review flags. "
        "This layer’s job is human factors — make the system usable without changing agent contracts.",
    )
    para(
        doc,
        "Orchestration layer: LangGraph StateGraph. This layer’s job is control flow and shared state. "
        "It decides the sequence of clinical and retrieval work and records an audit trail in agent_trace.",
    )
    para(
        doc,
        "Model and RAG layer: embeddings, vector search, optional LLM, extractive fallback, verification heuristics. "
        "This layer’s job is grounding — answers must be anchored in retrieved evidence.",
    )
    para(
        doc,
        "Data and platform layer today: seed JSON KB, chunks.json, faiss.index, source_registry.json, localhost runtime. "
        "This layer’s job is persistence and reproducibility for a short project under cost and install constraints.",
    )
    para(
        doc,
        "If something fails in a demo, this layer model tells you where to look. If retrieval is wrong, inspect vectorstore and evidence_rank. "
        "If citations are invalid, inspect verification. If the UI is confusing, change presentation only. "
        "That separation is how you keep a multi-week prototype maintainable.",
    )

    # ========== FIGURE 1 ==========
    h1(doc, "Part III — Figure 1: End-to-end system architecture")
    para(
        doc,
        "Figure 1 answers a simple question: when a clinician submits a case, what subsystems activate, in what order, and why?",
    )
    para(
        doc,
        "Start at Presentation. Streamlit is the primary surface because it is the fastest path to a credible interactive demo in Python "
        "without standing up a separate frontend build pipeline. That choice optimizes time-to-demo and grader reproducibility. "
        "The alternative — React or Next.js — would look more “production,” but it would add routing, auth scaffolding, and deployment complexity "
        "that does not increase RAG correctness for Milestone 2. Optional FastAPI /ask exists because some evaluators and future integrations "
        "prefer an HTTP contract; it calls the same pipeline, so we do not fork business logic.",
    )
    para(
        doc,
        "Next, Orchestration. LangGraph holds GraphState and compiles a deterministic graph. I chose LangGraph over a single procedural script "
        "because agentic CDSS requires visible step boundaries: assessment, risk, query, retrieval, verification, recommendation, transparency. "
        "I could have written one long Python function. That would be shorter. It would also hide the multi-agent structure the course asks us to demonstrate, "
        "and it would make agent_trace awkward. LangGraph’s cost is more code and an extra dependency; its benefit is clarity, debuggability, and alignment with agentic AI learning goals.",
    )
    para(
        doc,
        "Intelligence sits above retrieval: optional query rewrite with an LLM, verification that checks citation integrity and conflicts, "
        "and a recommendation builder that turns ranked hits plus patient/risk context into clinician-facing structure. "
        "Notice the design: intelligence is not allowed to invent a knowledge base. It consumes context that retrieval already grounded.",
    )
    para(
        doc,
        "Retrieval is FAISS plus TF-IDF plus priority re-rank. Why FAISS? It is a mature, local, low-ops vector index suitable for a small curated corpus. "
        "Why not Chroma or a managed vector database? Those are excellent green-cell options for multi-user scale and ops features, "
        "but they add service management. For nineteen documents and a class demo, FAISS on disk is proportionate. "
        "Why TF-IDF default rather than sentence-transformers? Because heavy neural embedding stacks broke in our environment — including dependency fragility around torch-related packages. "
        "TF-IDF is deterministic, offline, and cheap. The quality trade-off is real: neural embeddings usually improve semantic recall. "
        "Milestone 2 prefers a reliable demo over a brittle higher-ceiling stack. That is an engineering decision, not a claim that TF-IDF is state of the art forever.",
    )
    para(
        doc,
        "Knowledge is the seed medical KB and persisted artifacts. Curated public-style summaries let us control demo content and licensing posture: no PHI, no hospital EHR dump. "
        "The trade-off versus live crawl is freshness versus controllability. We mitigate freshness with a Knowledge Update probe loop, which I will explain later, "
        "while being honest that full automatic guideline PDF ingestion is still green-cell work.",
    )
    para(
        doc,
        "Put together, the request path is: UI captures patient and question → LangGraph nodes execute → retrieval and ranking produce evidence → "
        "verification and recommendation consume that evidence → transparency packages confidence, sources, and disclaimer for the clinician.",
    )

    # ========== FIGURE 2 ==========
    h1(doc, "Part IV — Figure 2: LangGraph execution graph and component communication")
    para(
        doc,
        "Figure 2 is the compiled graph in src/graph.py. This is the heart of agent communication.",
    )
    para(
        doc,
        "Agents do not call each other through ad-hoc imports in random order. They communicate through GraphState. "
        "Each node function receives state, writes the fields it owns, appends to agent_trace, and returns. "
        "The next edge decides who runs next. That is the interaction model: shared typed state plus linear edges for this prototype.",
    )
    para(
        doc,
        "Walk the edges with me. Entry is knowledge_update_agent. If run_knowledge_update is false, it records a skip and passes through. "
        "If true, it runs the URL probe workflow and can invalidate the in-memory store. Then patient_assessment_agent merges form input with text-extracted clues. "
        "risk_analysis_agent derives hypertension stage, obesity status, a prototype CV band, and preventive needs. "
        "query_agent builds a retrieval query, optionally LLM-rewritten, incorporating patient context bits so retrieval is not question-only. "
        "retrieval_agent searches FAISS with tier re-rank and builds the evidence block. "
        "verification_agent drafts a grounded summary, detects conflict themes, validates citation indices, computes confidence, and sets needs_human_review. "
        "recommendation_agent structures clinical recommendations and evidence summary. "
        "transparency_agent finalizes the clinician-facing answer with confidence, review recommendation, hierarchy policy, and disclaimer. Then END.",
    )
    para(
        doc,
        "Why linear edges rather than a complex branching supervisor? Because for Milestone 2, the clinical workflow is intentionally sequential: "
        "you assess before you risk-stratify, you retrieve before you recommend, you verify before you claim confidence. "
        "A more elaborate router could skip steps when data is missing; we instead keep the path stable and encode uncertainty into confidence and review flags. "
        "That favors reliability and demo predictability over adaptive control-flow sophistication.",
    )
    para(
        doc,
        "Latency implication: each node adds overhead, but for this corpus retrieval dominates or is still small locally. "
        "LLM nodes, when enabled, dominate wall time. That is why optional LLM plus extractive fallback is a latency and reliability control knob.",
    )

    # ========== FIGURE 3 ==========
    h1(doc, "Part V — Figure 3: Clinical multi-agent CDSS workflow")
    para(
        doc,
        "Figure 3 translates graph nodes into the clinical CDSS language the course expects. This is how I explain the system to a professor who cares about both medicine and engineering.",
    )
    para(
        doc,
        "Step 1, Patient Assessment, exists because retrieval without patient context answers a different question than the one in clinic. "
        "Age, sex, blood pressure, diabetes status, LDL, BMI, smoking, ASCVD history, and other CV risks change which guidelines apply. "
        "We accept form input and also extract clues from free text so the demo still works if someone only types a narrative question.",
    )
    para(
        doc,
        "Step 2, Risk Analysis, turns raw fields into intermediate clinical structure: hypertension stage, obesity status, a prototype cardiovascular risk band, preventive needs. "
        "This is not a validated clinical calculator claiming FDA-grade risk scores. It is a transparent prototype layer that conditions query formulation and recommendation framing. "
        "Engineering-wise, putting risk in its own node makes the logic testable and visible in the agent trace.",
    )
    para(
        doc,
        "Step 3, Evidence Retrieval, is where RAG meets evidence-based medicine. We do not treat the highest cosine neighbor as the highest clinical authority. "
        "Priority ranking enforces guidelines first when relevant.",
    )
    para(
        doc,
        "Step 4, Evidence Verification, is the anti-hallucination and anti-overconfidence layer. We check that cited indices exist, look for cross-source conflict themes, "
        "and reduce confidence when patient data is sparse, when top evidence is low-tier, or when superseded material appears. "
        "needs_human_review is a first-class output, not an afterthought.",
    )
    para(
        doc,
        "Steps 5 and 6 produce recommendations and an evidence summary with organization, year, and citations. "
        "Step 7, Transparency, is the trust UI contract: show your work. Confidence, sources consulted, hierarchy policy, agent trace, disclaimer.",
    )
    para(
        doc,
        "From a systems view, this is human-in-the-loop design. The machine proposes; the clinician disposes. "
        "That is both an ethical stance and an engineering stance: we optimize for grounded support, not autonomous treatment.",
    )

    # ========== FIGURE 4 ==========
    h1(doc, "Part VI — Figure 4: RAG pipeline, indexing and query time")
    para(
        doc,
        "Retrieval-augmented generation has two clocks: offline indexing and online answering. Confusing them is how students accidentally claim live crawl when they only have a static index.",
    )
    para(
        doc,
        "Offline path — scripts/build_index.py. Load seed documents. Clean text. Chunk with overlap so a guideline statement is less likely to be split awkwardly. "
        "Fit TF-IDF on chunk text and produce vectors. Create FAISS IndexFlatIP and add vectors. Persist faiss.index, chunks.json, tfidf.pkl, and meta.json. "
        "IndexFlatIP means we use inner product over normalized or compatible vectors for similarity; for this prototype it is an exact search index, which is appropriate at small N. "
        "At our scale — roughly nineteen documents and twenty chunks — exact search latency is negligible. Approximations like HNSW would be a scalability optimization for much larger corpora.",
    )
    para(
        doc,
        "Online path — retrieval_agent. The query agent may already have appended patient context. We embed the query with the same embedder, search FAISS for candidates, "
        "re-rank by evidence tier and quality, filter weak hits, prefer non-superseded guidelines, and format a numbered evidence block.",
    )
    para(
        doc,
        "Generation path — verification and recommendation. If llm_configured(), we ask the chat model to summarize guideline-concordant points with [n] citations and to call out conflicts. "
        "If the LLM errors or no key exists, extractive_answer builds a grounded summary from chunks. Then citation regex validation runs. "
        "Unsupported citation indices reduce trust and feed review flags.",
    )
    para(
        doc,
        "Why this RAG design instead of “just prompt GPT with medical advice”? Because the failure mode of plain LLMs in medicine is confident fabrication of guidelines and doses. "
        "RAG does not eliminate error, but it changes the error profile toward retrieval mistakes you can inspect, rather than invisible invention. "
        "That inspectability is the point of source-linked CDSS.",
    )
    para(
        doc,
        "Cost: TF-IDF and local FAISS are effectively free at demo scale. Ideal neural embeddings and re-rankers improve quality and cost money or GPU time. "
        "Latency: local retrieval is fast; LLM calls dominate when enabled. Reliability: offline extractive path means the demo survives without network or API quota.",
    )

    # ========== FIGURE 5 ==========
    h1(doc, "Part VII — Figure 5: Evidence hierarchy as a clinical ranking prior")
    para(
        doc,
        "Figure 5 is the policy encoded in evidence_rank.py. Think of it as a clinical prior placed on top of vector search.",
    )
    para(
        doc,
        "Priority 1: Clinical Practice Guidelines — USPSTF, ADA, ACC/AHA, NCCN, NICE, WHO, and related societies. "
        "Priority 2: Systematic reviews and meta-analyses. Priority 3: Randomized trials. Priority 4: PubMed/PMC-style literature. "
        "Priority 5: AHRQ. Priority 6: MedlinePlus patient education. Priority 7: MedQuAD for testing only.",
    )
    para(
        doc,
        "Why hard sort by priority rather than a single blended score? Because early experiments with pure semantic ranking overweighted education-style pages. "
        "Those pages can be textually similar to a question while being the wrong authority class for clinical decision support. "
        "A blended score can hide that mistake inside a float. A hard tier makes the policy auditable: you can look at the top hit and see its priority.",
    )
    para(
        doc,
        "Within a tier, quality signals still matter — evidence level, recommendation strength, recency where appropriate — and superseded documents are downgraded. "
        "That addresses another real failure mode: an outdated guideline version that still embeds well.",
    )
    para(
        doc,
        "This hierarchy also connects Milestone 1 to Milestone 2. In Milestone 1 we argued that guidelines are usually the best starting point when relevant. "
        "Milestone 2 implements that argument as code, not as a slide slogan.",
    )

    # ========== FIGURE 6 ==========
    h1(doc, "Part VIII — Figure 6: Retrieval and ranking pipeline step by step")
    para(
        doc,
        "Now the data flow inside one retrieval call, step by step, as if we are tracing a single ask.",
    )
    para(
        doc,
        "Step 1 — Embed query. The same TF-IDF space as the index. If spaces diverge, retrieval collapses. That is why embedder choice is locked between build_index and runtime load.",
    )
    para(
        doc,
        "Step 2 — Candidate search. FAISS returns a wider candidate_k than final top-k. Why? Because clinical re-ranking needs room to promote a guideline that was not the top semantic neighbor. "
        "Wide recall, precise reorder — a classic IR pattern.",
    )
    para(
        doc,
        "Step 3 — Annotate tiers. Infer Priority 1–7 from source type, organization, and metadata. Missing metadata is annotated rather than silently treated as equal.",
    )
    para(
        doc,
        "Step 4 — Re-rank. Sort by tier ascending — one is best — then quality within tier; penalize superseded.",
    )
    para(
        doc,
        "Step 5 — Filter and cut. Drop extremely weak semantic matches when appropriate, keep top-k, prefer active guidelines in the final window.",
    )
    para(
        doc,
        "Step 6 — Context pack. format_evidence_block creates the [1]…[k] structure used by generation and citation checking. "
        "This is the interface contract between retrieval and language generation. If that contract is loose, citation verification becomes theater.",
    )
    para(
        doc,
        "Compared with alternatives: a cross-encoder re-ranker would likely improve ordering further — that is a green-cell upgrade. "
        "Hybrid BM25 plus dense retrieval is another green upgrade for robustness. We did not implement them in Milestone 2 because TF-IDF plus tier re-rank already corrected the main clinical failure mode we observed, "
        "and because added models increase dependency and latency risk for the demo.",
    )

    # ========== FIGURE 7 ==========
    h1(doc, "Part IX — Figure 7: Module interaction map")
    para(
        doc,
        "Figure 7 is the repository as a runtime graph. Professors often ask: where does the code live, and who calls whom?",
    )
    para(
        doc,
        "app.py is the Streamlit process. It collects patient_input and question, then invokes the pipeline. It also hosts maintenance buttons that call knowledge update and rebuild.",
    )
    para(
        doc,
        "src/graph.py is the orchestrator. It imports CDSS helpers, LLM helpers, and get_store for retrieval. It owns sequencing.",
    )
    para(
        doc,
        "src/cdss.py owns patient merge, risk analysis, conflict detection helpers, and recommendation structuring. "
        "Keeping clinical structuring outside the graph file prevents graph.py from becoming an unmaintainable monolith.",
    )
    para(
        doc,
        "src/llm.py isolates model access: chat_json_or_text, extractive_answer, evidence formatting. "
        "If the provider changes — OpenAI, OpenRouter, Ollama via base URL — only this boundary should move.",
    )
    para(
        doc,
        "src/vectorstore.py owns FAISS build/load/search. src/embeddings.py owns TF-IDF and optional API embeddings. "
        "src/evidence_rank.py owns tier inference and re-rank. src/knowledge_update.py owns registry probes and supersession. "
        "api/main.py optionally wraps the same invoke path for HTTP.",
    )
    para(
        doc,
        "The communication pattern is therefore hub-and-spoke through GraphState, not a mesh of hidden side effects. "
        "That is maintainability: a new UI does not rewrite retrieval; a new embedder does not rewrite Streamlit forms.",
    )

    # ========== FIGURE 8 ==========
    h1(doc, "Part X — Figure 8: Knowledge update workflow")
    para(
        doc,
        "Clinical knowledge drifts. A CDSS that never updates is a museum. But automatic ingest of every PDF from every society is a large systems problem — legal, parsing, versioning, governance.",
    )
    para(
        doc,
        "Our implemented prototype is deliberately scoped. source_registry.json lists authoritative URLs. "
        "We HTTP-probe them for availability and headers such as Last-Modified or ETag when present. "
        "We compare against local seed document expectations, apply local superseded_by rules, optionally rebuild FAISS, and log a report.",
    )
    para(
        doc,
        "Triggers: Streamlit sidebar maintenance actions, or the graph flag when enabled. "
        "Organizations probed in the prototype include USPSTF, ADA, ACC/AHA-related pages, MedlinePlus, AHRQ, and PMC-related sources in the registry design.",
    )
    para(
        doc,
        "What this gives us: an executable freshness loop for the demo and a clear place to hang future ETL. "
        "What this does not give us yet: automatic download and chunking of every new NCCN or KDIGO PDF into the vector store with clinician approval workflows. "
        "That green path matters for production, but claiming it now would violate the honesty contract.",
    )

    # ========== FIGURE 9 ==========
    h1(doc, "Part XI — Figure 9: Deployment today versus production target")
    para(
        doc,
        "Deployment is where students most often over-claim. So I separate yellow and green explicitly.",
    )
    para(
        doc,
        "Today’s yellow deployment is local: laptop or classroom machine, Python virtual environment, streamlit run app.py, optional uvicorn for FastAPI, "
        "indexes on disk, optional API key in .env, public seed KB, no PHI. "
        "Why is that appropriate for Milestone 2? Because the grading event needs reproducibility, near-zero cost, and resilience when campus networks or API quotas fail. "
        "Local-first maximizes reliability for a timed demo.",
    )
    para(
        doc,
        "Ideal green deployment looks like a real healthcare AI service: HIPAA-eligible cloud on Azure or AWS, managed vector search with hybrid retrieval, "
        "enterprise LLM endpoints such as Azure OpenAI, Bedrock, or Claude API under contract, a React clinician console with authentication, "
        "CI evaluation gates, observability, guideline watchers with human approval, and only then optional EHR/FHIR under a BAA.",
    )
    para(
        doc,
        "Cost jumps from near zero to recurring cloud and model spend. Complexity jumps from one repo to platform engineering. "
        "Scalability and multi-tenant safety become possible. Maintainability requires ownership. Latency needs caching and SLOs. "
        "Reliability becomes on-call, not “restart Streamlit.” That is why green is better for production — and why it is not Milestone 2’s deliverable.",
    )

    # ========== TECH STACK ==========
    h1(doc, "Part XII — AI Tech Stack table: answering the instructor’s three questions")
    para(
        doc,
        "This slide follows AI Tech Stack.docx. The legend is intentional and visual: yellow means implemented in Milestone 2; green means ideal future architecture. "
        "I will answer the three required questions in conference form.",
    )

    h2(doc, "Question 1 — What stack did we implement?")
    para(
        doc,
        "Data layer and vector store: local JSON medical KB and FAISS local index. "
        "I selected JSON plus FAISS because the corpus is curated and small, and because local files are easy to grade and version in a student repo. "
        "Alternatives like SQLite or Supabase would help multi-user writes; BigQuery or Snowflake would help analytics warehouses; Azure AI Search or OpenSearch would help managed hybrid retrieval. "
        "Those are green because they optimize scale and ops, not Milestone 2 demo certainty.",
    )
    para(
        doc,
        "Model training and MLOps: VS Code and local Python scripts. We are not training a foundation model; we are assembling a RAG system. "
        "Colab, W&B, Vertex, SageMaker, Azure ML, Databricks are green when you need experiment tracking and training pipelines at scale.",
    )
    para(
        doc,
        "Coding, IDE, and agents: terminal scripts, VS Code, Cursor, and LangGraph with LangChain ecosystem components. "
        "LangGraph was selected specifically to implement an explicit multi-agent CDSS graph rather than a hidden script. "
        "CrewAI or N8N can be green alternatives for different orchestration styles; GitHub Copilot is green as an assistive coding accelerator, not as runtime CDSS logic.",
    )
    para(
        doc,
        "GenAI models: Ollama-compatible local path and OpenAI-compatible API via environment configuration. "
        "This dual path is deliberate. Local or compatible endpoints support privacy-conscious and offline demos; hosted APIs improve fluency when available. "
        "Gemini, Bedrock, Azure OpenAI, Anthropic Claude are green production options with different enterprise and clinical contracting profiles.",
    )
    para(
        doc,
        "Frontend: Streamlit implemented. HTML prototypes are too weak for the interactive CDSS we need; React/Next is the green product UI. "
        "Hosting: localhost Streamlit and Uvicorn. Render, Fly, Cloud Run, Lambda, Azure App Service, and HIPAA Kubernetes are green runtime targets.",
    )
    para(
        doc,
        "Bonus implemented rows that matter for this project: scikit-learn TF-IDF plus FAISS top-k; LangGraph StateGraph CDSS; local eval JSON for latency and citations. "
        "Those bonus rows are where the course themes — RAG, agents, evaluation — actually live.",
    )

    h2(doc, "Question 2 — What is the ideal stack?")
    para(
        doc,
        "Ideal is the green column: managed clinical knowledge storage, cloud MLOps, stronger hosted LLMs, React clinician applications, HIPAA-eligible hosting, "
        "neural embeddings with hybrid search and cross-encoder or vendor re-rankers, managed workflow orchestration where appropriate, "
        "and continuous observability with groundedness monitoring and clinician override analytics. "
        "That stack is what you build when the user base, liability, and uptime requirements become real.",
    )

    h2(doc, "Question 3 — Why those ideal components? Trade-offs")
    para(
        doc,
        "I will not treat trade-offs as a slide footnote. They are the reason yellow and green differ.",
    )
    para(
        doc,
        "Cost: Yellow keeps the project near zero recurring cost — critical for a university prototype. "
        "Green improves quality with paid embeddings, LLMs, and cloud, but introduces monthly spend and, for PHI, legal cost around BAAs.",
    )
    para(
        doc,
        "Complexity: Yellow is one Python repository a grader can run. Green multiplies moving parts — frontend, identity, managed search, secrets, networking. "
        "Complexity is not automatically sophistication; unmanaged complexity is risk.",
    )
    para(
        doc,
        "Scalability: Yellow FAISS-on-disk is correct for tens of documents. It is the wrong long-term answer for multi-clinic concurrent retrieval over large corpora. "
        "Green managed vector and autoscaling APIs exist for that future load.",
    )
    para(
        doc,
        "Maintainability: Yellow’s LangGraph boundaries and seed JSON are easy to inspect in a semester. "
        "Green needs CI, evaluation gates, and ownership so guideline drift does not silently poison answers.",
    )
    para(
        doc,
        "Latency: Yellow can be very fast locally, especially without LLM. Green’s neural re-rankers and always-on LLMs can improve answer quality while adding milliseconds to seconds; "
        "caching, async jobs, and streaming become necessary design elements.",
    )
    para(
        doc,
        "Reliability: Yellow’s offline extractive path is a reliability feature for demos. Green reliability is different — redundancy, monitoring, rollback, incident response. "
        "A system can be locally reliable and still be production-unready.",
    )
    para(
        doc,
        "Safety and trust: Yellow implements citations, confidence, review flags, and disclaimers. "
        "Green adds clinical validation programs, audit logging, red-teaming, and governance — the difference between a course prototype and a system you would let near real care workflows.",
    )

    # ========== DATA INSIGHTS ==========
    h1(doc, "Part XIII — Data, model behavior, and insights from building it")
    para(
        doc,
        "The data story must stay precise. We use a public-style seed knowledge base: guideline summaries, MedlinePlus-style pages, AHRQ themes, PubMed/PMC-style notes, and a MedQuAD sample for testing. "
        "No private patient records. The pipeline is load, clean, chunk, embed, store. Current index metadata is on the order of 19 documents and 20 chunks with embedder local:tfidf. "
        "Optional PubMed E-utilities helper exists; bulk PMC crawl and EHR/FHIR are not claimed.",
    )
    para(
        doc,
        "Model behavior: with an API key, chat completion can improve fluency while remaining citation-bound. Without it, extractive mode still produces a grounded demo. "
        "That dual-mode design is how we satisfy both “LLMs” and “works offline” course pressures.",
    )
    para(
        doc,
        "Insights that changed the architecture: guideline-first ranking beat pure semantic retrieval for trust. Superseded docs must be penalized. "
        "Citation checking must be code, not hope. Offline path is mandatory for reliable classroom demos. "
        "And CDSS step structure is more clinically usable than a single chatbot box — because clinicians think in assessment, evidence, and recommendation, not only in chat turns.",
    )

    # ========== EVAL ==========
    h1(doc, "Part XIV — Evaluation metrics and why they matter")
    para(
        doc,
        "In AI engineering, if you cannot measure it, you cannot defend it. Our evaluation script maps proposal metrics onto what this prototype can honestly measure now.",
    )
    para(
        doc,
        "Citation precision on the demo evaluation set reached 1.0 — every [n] citation referred to a real retrieved hit. "
        "Why that metric matters: it is a direct check on groundedness plumbing. A fluent answer with fabricated citation indices is worse than a blunt extractive answer with valid citations.",
    )
    para(
        doc,
        "Chunks per query at five after priority re-rank shows the system is retrieving a usable evidence set rather than returning empty context. "
        "Corpus size — 19 documents, 20 chunks — keeps expectations honest: this is a controlled prototype corpus, not a claim of indexing all of PubMed.",
    )
    para(
        doc,
        "The qualitative RAG versus plain LLM comparison is essential. With RAG: source links, constrained context, agent-trace audit, offline capability. "
        "Plain LLM: no retrieval audit trail, hallucination risk on guidelines, and often unavailable without an API key in our demo setting. "
        "That comparison is the proposal’s technical evaluation idea, implemented as an engineering argument rather than a marketing claim.",
    )
    para(
        doc,
        "What we do not over-claim: we do not present multi-site clinician user-satisfaction RCTs. Business metrics like time saved are future study designs. "
        "Milestone 2 prioritizes technical groundedness metrics that the repository can reproduce.",
    )

    # ========== CHALLENGES ==========
    h1(doc, "Part XV — Challenges, learnings, and engineering judgment")
    para(
        doc,
        "The hardest lessons were operational. Neural embedding stacks failed installs; we chose TF-IDF to protect the demo. "
        "Semantic search alone elevated the wrong authority class; we imposed evidence hierarchy. "
        "Superseded guidelines still retrieved; we added penalties. UI readability required CSS and layout work without rewriting agents — "
        "a reminder that presentation and orchestration should stay decoupled. "
        "And medical AI requires loud uncertainty: disclaimers and human-review flags are part of the system, not cover slides.",
    )
    para(
        doc,
        "The learning I want the audience to remember: separating implemented and ideal is not humility theater. It is how AI engineers communicate risk, scope, and next investment.",
    )

    # ========== DEMO ==========
    h1(doc, "Part XVI — Live demo narrative (speak this while the application runs)")
    para(
        doc,
        "I will now switch to the running system. Camera remains on. The codebase path is AI HEALTH CARE.",
        italic=True,
    )
    para(
        doc,
        "First I confirm the index is built — python scripts\\build_index.py if needed — then launch Streamlit. "
        "You should see the Clinical Evidence CDSS workspace: navy-to-blue header, sidebar with system status KPIs, and the patient assessment card.",
    )
    para(
        doc,
        "I enter a realistic outpatient-style vignette. For example: an adult with clinical ASCVD questions about statin intensity, "
        "or a type 2 diabetes question about A1C targets, with age, blood pressure, diabetes status, LDL, BMI, and smoking filled when known. "
        "Leaving fields unknown is allowed; the system should become more conservative via confidence and review flags when data is sparse.",
    )
    para(
        doc,
        "I submit the clinical query and walk the outputs in clinical order. Patient assessment — did we capture the fields? "
        "Risk analysis — what intermediate structure was inferred? Evidence tiles — are guidelines appearing first when they should? "
        "Verification — any conflicts, what confidence, is human review recommended? Recommendations — are citations attached? "
        "Transparency — can I see agent_trace and sources consulted?",
    )
    para(
        doc,
        "I briefly open maintenance: Knowledge Update probes and Rebuild index. I state their implemented scope — URL freshness and local supersession — "
        "and I explicitly say automatic full PDF ingest is future work.",
    )
    para(
        doc,
        "Finally I point to the code: graph.py for edges, evidence_rank.py for hierarchy, vectorstore.py for FAISS. "
        "The point of the demo is not animation; it is correspondence between slides and executable artifacts.",
    )
    para(
        doc,
        "I close the demo the same way I opened the talk: research prototype, not a medical device, clinician decides.",
    )

    # ========== CLOSING ==========
    h1(doc, "Closing")
    para(
        doc,
        "Let me leave you with the engineering thesis of this project.",
    )
    para(
        doc,
        "Better clinical decisions need evidence clinicians can verify. "
        "A language model without retrieval can sound certain while being uncheckable. "
        "Our Milestone 2 system demonstrates a different pattern: retrieve public medical knowledge, rank it by clinical priority, "
        "orchestrate CDSS steps with LangGraph agents, verify citations, expose confidence and review flags, and keep the human in the loop.",
    )
    para(
        doc,
        "We chose a yellow stack that optimizes cost, reliability, and demo reproducibility under semester constraints. "
        "We mapped a green stack that would be more appropriate for production scale, governance, and quality. "
        "We measured citation integrity and showed a working Streamlit path. "
        "That is AI engineering for Healthcare AI: not only models, but systems, trade-offs, and honest scope.",
    )
    para(
        doc,
        "I’m Audrey Rah, University of Houston. I’m happy to take questions on architecture, ranking policy, evaluation, or the live system.",
    )

    # ========== Q&A APPENDIX ==========
    h1(doc, "Appendix — Extended Q&A preparation (if asked)")
    qa = [
        (
            "Why not claim neural embeddings as implemented?",
            "Because the default runnable path is TF-IDF after neural stacks proved fragile in our environment. Claiming neural as default would misrepresent what graders can reproduce without extra setup.",
        ),
        (
            "Is MedQuAD used as clinical authority?",
            "No. It is Priority 7, testing only. Guidelines and higher evidence classes outrank it by policy.",
        ),
        (
            "Can this integrate with Epic or FHIR tomorrow?",
            "Not as implemented. EHR/FHIR under HIPAA controls is green-cell work requiring identity, BAAs, and clinical safety processes far beyond this prototype.",
        ),
        (
            "What happens when sources conflict?",
            "Verification detects conflict themes, confidence is reduced, and needs_human_review can trigger. The system should surface disagreement rather than silently average it away.",
        ),
        (
            "Why Streamlit rather than a “real” web app?",
            "Time-to-interactive-demo and Python-native integration with the RAG stack. React is better for polished multi-page product UX and is listed as ideal.",
        ),
        (
            "How long does retrieval take?",
            "On the local prototype corpus, retrieval is typically well under a second; LLM calls dominate when enabled. Exact numbers vary by machine; evaluate.py records timings for the run environment.",
        ),
    ]
    for q, a in qa:
        para(doc, f"Q: {q}", bold=True, after=2)
        para(doc, f"A: {a}", after=10)

    h1(doc, "Delivery timing guide (approximately 15–20 minutes)")
    for line in [
        "0:00–1:30 — Opening problem + Milestone 2 contract + disclaimer",
        "1:30–3:30 — What was built + AI engineering layers",
        "3:30–8:30 — Figures 1–4 (architecture, LangGraph, CDSS, RAG)",
        "8:30–11:30 — Figures 5–9 (hierarchy, ranking, modules, update, deployment)",
        "11:30–14:00 — Tech stack yellow/green + trade-offs",
        "14:00–15:30 — Data insights + evaluation + challenges",
        "15:30–18:30 — Live demo narrative",
        "18:30–20:00 — Closing + questions buffer",
    ]:
        para(doc, f"• {line}", after=4)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
