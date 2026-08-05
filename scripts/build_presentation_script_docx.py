#!/usr/bin/env python3
"""Build Milestone 2 complete presentation speaker script (.docx)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"C:\Users\audre\OneDrive\1- Final project\AI HEALTH CARE\Milestone2_Presentation_Script.docx")


def set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=RGBColor(0x0F, 0x4C, 0x81))
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_slide(doc, number: int, title: str, timing: str, body: str, transition: str = ""):
    add_heading_styled(doc, f"Slide {number} — {title}", level=1)
    add_para(doc, f"Suggested timing: {timing}", size=10, italic=True, space_after=6)
    add_para(doc, "SPEAKER NOTES", size=11, bold=True, space_after=4)
    for para in body.strip().split("\n\n"):
        add_para(doc, para.strip(), size=11, space_after=8)
    if transition:
        add_para(doc, "TRANSITION", size=11, bold=True, space_after=4)
        add_para(doc, transition.strip(), size=11, italic=True, space_after=14)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Milestone 2 Presentation Script")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x0F, 0x4C, 0x81))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Source-Linked Clinical Evidence CDSS")
    set_run_font(r2, size=16, bold=True)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "Complete speaker notes for recording and live delivery\n"
        "Audrey Rah · University of Houston · Houston, Texas, USA\n"
        "Codebase: AI HEALTH CARE/"
    )
    set_run_font(r3, size=11)

    add_para(
        doc,
        "How to use this script: Speak naturally — do not read every sentence word-for-word. "
        "Use each slide’s notes as a talk track. Bold ideas are the points graders care about: "
        "what is implemented, what is ideal, and why the engineering trade-offs matter. "
        "Always state clearly: this is a research/education prototype, not a medical device.",
        size=11,
        italic=True,
    )

    add_heading_styled(doc, "Presentation goals (say this once early)", level=2)
    add_para(
        doc,
        "By the end of this talk, the audience should understand: (1) I built an end-to-end, "
        "runnable clinical evidence CDSS with RAG and LangGraph agents; (2) evidence is ranked "
        "by clinical priority, not semantic similarity alone; (3) yellow stack components are "
        "implemented today and green components are the ideal future architecture; (4) evaluation "
        "and a live demo prove the system works.",
    )

    # ------------------------------------------------------------------
    slides = [
        (
            1,
            "Title",
            "~20–30 seconds",
            """
Hi everyone — I’m Audrey Rah, University of Houston. This is my Milestone 2 presentation for the Source-Linked Clinical Evidence CDSS.

Milestone 1 was about the clinical problem and the persona. Milestone 2 is the engineering proof: a working end-to-end prototype that combines retrieval-augmented generation, LangGraph multi-agent orchestration, evidence-hierarchy ranking, and a clinician-facing Streamlit workspace.

The codebase lives in the AI HEALTH CARE folder. Before I go further: this is a research and education prototype. It is not a medical device, and a licensed clinician always makes the final decision.
""",
            "I’ll start with what was actually built, then walk the architecture figures, the AI tech stack, trade-offs, evaluation, and a live demo.",
        ),
        (
            2,
            "Agenda",
            "~30–40 seconds",
            """
Here’s the path for this talk.

First, I’ll summarize what was built — the honesty boundary between implemented and proposed.

Second, architecture: end-to-end layers, the LangGraph execution graph, RAG, ranking, and data flow.

Third, the clinical multi-agent CDSS workflow and the knowledge-update loop.

Fourth, the instructor AI Tech Stack table — yellow for implemented, green for ideal — and the trade-offs behind those choices.

Fifth, data, insights, evaluation metrics, challenges, and then the live demo that proves the system runs.
""",
            "Starting with the resume-style summary of what exists in code today.",
        ),
        (
            3,
            "What Was Built (Implemented — Not Proposed)",
            "~1.5–2 minutes",
            """
This slide is intentionally titled “Implemented — Not Proposed,” because Milestone 2 is judged on a working system.

First: I built an end-to-end clinical decision-support prototype. The clinician UI is Streamlit. Behind it is a LangGraph StateGraph with eight nodes. The system produces source-linked recommendations with citations, a confidence score, and a human-review flag when evidence or patient data is weak.

Second: hierarchy-aware RAG over a public-style medical knowledge base. Vectors live in FAISS using IndexFlatIP. Embeddings default to local TF-IDF so the demo runs offline. Retrieved chunks are re-ranked by Priority 1 through 7 — guidelines first, MedQuAD last — and superseded guidelines are downgraded. If an OpenAI-compatible API key is present, the system can use a chat LLM; if not, it falls back to an extractive grounded summary from retrieved chunks. That offline path is a deliberate reliability decision.

Third: explainability and measurement. The UI shows evidence tiles, conflict notes, agent traces, and metrics. There is an evaluation script writing citation precision and latency into JSON. FastAPI /ask is optional. A Knowledge Update Agent can probe authoritative URLs and apply local supersession.

Every claim on this slide maps to runnable code in AI HEALTH CARE — not a deck-only architecture.
""",
            "That honesty boundary is the AI engineering lens for the whole talk — Chip Huyen–style: ship under constraints, and don’t blur prototype with production.",
        ),
        (
            4,
            "AI Engineering Lens (Chip Huyen–style)",
            "~1–1.5 minutes",
            """
I organize the system the way production AI systems are usually reasoned about — by layers.

Application layer: Streamlit clinician workspace, patient form, clinical query, citations, transparency panel, disclaimer, and review flags.

Orchestration layer: LangGraph StateGraph, shared GraphState, deterministic edges between agents, and an agent_trace for auditability.

Model and RAG layer: optional chat LLM, extractive fallback, TF-IDF embedding into FAISS retrieval, then tier re-rank and verification.

Data and platform layer today: seed JSON medical KB, chunk store plus FAISS index on disk, source registry for URL probes, and localhost runtime.

The principle for this presentation: yellow cells are runnable today; green cells are the production-oriented roadmap. I will never present green items as if they already ship. That distinction is itself part of the Milestone 2 deliverable.
""",
            "With that framing, let’s look at Figure 1 — how a request moves end-to-end through the implemented path.",
        ),
        (
            5,
            "Figure 1 — End-to-End System Architecture",
            "~1.5–2 minutes",
            """
Figure 1 shows what a clinician request actually touches.

At the presentation layer: Streamlit CDSS UI — that’s the main demo surface — and optionally FastAPI /ask if we want an HTTP interface.

Orchestration: LangGraph CDSS StateGraph with GraphState and agent_trace. This is the control plane. It decides the sequence of clinical and retrieval steps and carries intermediate results.

Intelligence: query rewrite when an LLM is configured, verification with confidence, and the recommendation builder that turns ranked evidence into clinician-facing structure.

Retrieval: TF-IDF embedder, FAISS IndexFlatIP, and Priority 1–7 re-ranking. This is where “RAG” becomes concrete — retrieve first, then generate or extract grounded text.

Knowledge: seed_medical_kb.json, persisted chunks.json and faiss.index, and source_registry.json for update probes.

Request flow in one sentence: UI → LangGraph nodes → retrieve and rank → verify → recommend → transparency with citations.

Why this architecture? Separating presentation, orchestration, retrieval, and knowledge makes the system demoable, debuggable, and honest about which layer failed when something looks wrong.
""",
            "Figure 2 zooms into the LangGraph edges that implement that flow.",
        ),
        (
            6,
            "Figure 2 — LangGraph Execution Graph",
            "~1.5–2 minutes",
            """
This figure is not conceptual — it matches src/graph.py.

Entry point is knowledge_update_agent. Then edges go linearly: patient assessment → risk analysis → query agent → evidence retrieval → evidence verification → recommendation plus evidence summary → transparency and final answer → END.

Knowledge Update is optional per invoke via a run_knowledge_update flag. In the UI, that corresponds to maintenance actions; for a normal ask, it usually skips. All other nodes always execute.

Shared GraphState carries the question, patient input, assessment, risk analysis, cleaned query, hits, verification, recommendation, timings, and agent_trace.

Why LangGraph instead of one giant function? Clear agent boundaries make the CDSS steps visible in the demo, produce an audit trail, and match the course emphasis on agentic AI. The trade-off is slightly more code than a single RAG call — and that trade-off is intentional for clarity and grading.
""",
            "Figure 3 maps those graph nodes onto the clinical CDSS steps the course asked for.",
        ),
        (
            7,
            "Figure 3 — Multi-Agent CDSS Workflow",
            "~1.5–2 minutes",
            """
This is the clinician-facing story of the same graph.

Step 1 — Patient Assessment: form fields plus text extraction for age, blood pressure, diabetes, LDL, BMI, smoking, ASCVD, and related risks.

Step 2 — Risk Analysis: hypertension stage, obesity status, a prototype CV risk band, and preventive-care needs derived from that patient picture.

Step 3 — Evidence Retrieval: FAISS top-k with Priority 1–7 hierarchy, preferring active guideline versions.

Step 4 — Evidence Verification: conflict themes across sources, citation index checks, confidence score, and whether human review is recommended.

Steps 5–6 — Recommendation and Evidence Summary: lifestyle, medication, screening, follow-up style recommendations tied back to organizations, years, and citations.

Step 7 — Transparency: confidence, sources consulted, hierarchy policy, agent trace, and the disclaimer.

Persona continuity from Milestone 1: think of Dr. Lin, a resident who needs source-linked evidence between patients. The system supports judgment; it does not replace it.
""",
            "Next I’ll explain the RAG pipeline that powers Step 3 — indexing offline and retrieving at query time.",
        ),
        (
            8,
            "Figure 4 — RAG Pipeline",
            "~1.5–2 minutes",
            """
This is classic RAG with clinical controls — not generation-only.

Offline indexing path, scripts/build_index.py: load seed KB JSON → clean and chunk → fit TF-IDF and embed → build FAISS IndexFlatIP → persist index and metadata. Today that yields on the order of 19 documents and 20 chunks with embedder local:tfidf.

Online query path inside retrieval_agent: take the clinical query — often enriched with patient context — embed it in the same vector space, run inner-product search, then apply tier re-ranking and package a numbered context block.

Grounded generation and verification: retrieved context with [n] markers goes to an LLM if configured, otherwise extractive fallback. Citation indices are checked. Confidence and review flags are computed. Final answer plus sources are returned.

Why RAG here? A plain LLM can invent guidelines. Retrieval-first plus citation checks reduce unsupported claims. That’s the core safety and trust argument for this project.
""",
            "The next figure is the clinical reason semantic similarity alone is not enough — the evidence hierarchy.",
        ),
        (
            9,
            "Figure 5 — Evidence Hierarchy (Priority 1 → 7)",
            "~1.5 minutes",
            """
Implemented in src/evidence_rank.py. Sorting is hard by clinical priority, then by quality within the same priority — not similarity-only, and not “newest always wins.”

Priority 1: Clinical Practice Guidelines — USPSTF, ADA, ACC/AHA, and similar societies.

Priority 2: Systematic reviews and meta-analyses.

Priority 3: Randomized clinical trials.

Priority 4: PubMed / PubMed Central style literature.

Priority 5: AHRQ.

Priority 6: MedlinePlus — useful patient education, but labeled lower.

Priority 7: MedQuAD — testing and evaluation only, not a guideline substitute.

This hierarchy is how the system stays aligned with evidence-based medicine practice from Milestone 1: when a relevant guideline exists, prefer it.
""",
            "Figure 6 shows how that hierarchy is applied inside the retrieval pipeline at runtime.",
        ),
        (
            10,
            "Figure 6 — Retrieval & Ranking Pipeline",
            "~1.5 minutes",
            """
Hybrid retrieval in six steps.

One: embed the query with TF-IDF in the same space as the index.

Two: FAISS inner-product search over a wider candidate set for recall.

Three: annotate each hit with Priority 1–7, organization, year, evidence level.

Four: re-rank — higher tier wins; within tier, quality; superseded documents are penalized.

Five: apply score floors, cut to top-k, prefer non-superseded active guidelines.

Six: format_evidence_block so downstream agents see [1]…[k] context they can cite.

Engineering rationale: semantic search alone overweighted patient-education pages in early experiments. Clinical re-ranking fixes that failure mode without requiring a paid neural re-ranker in the Milestone 2 prototype.
""",
            "Figure 7 connects those algorithms to the actual Python modules in the repo.",
        ),
        (
            11,
            "Figure 7 — Component Interaction & Data Flow",
            "~1–1.5 minutes",
            """
At the center is LangGraph in src/graph.py. That is the orchestrator.

Around it: app.py Streamlit UI calls into the pipeline; cdss.py handles patient merge, risk analysis, and recommendation structure; llm.py provides chat or extractive answers; api/main.py optionally exposes FastAPI /ask.

Below: vectorstore.py for FAISS search; evidence_rank.py for tier re-rank; embeddings.py for TF-IDF; knowledge_update.py for URL probes.

Runtime contract is simple: invoke with question, patient_input, and run_knowledge_update, and get back a GraphState dictionary. That single entry point keeps UI, API, and CLI demos consistent.
""",
            "One of those modules deserves its own figure — knowledge update — because freshness is part of clinical trust.",
        ),
        (
            12,
            "Figure 8 — Knowledge Update Workflow",
            "~1.5 minutes",
            """
This is an implemented prototype loop — I want to be precise about scope.

Flow: source_registry.json → HTTP probe for status and optional Last-Modified or ETag → compare to local seed versions → apply local supersession rules → optionally rebuild FAISS → write knowledge_update_log.json.

It can be triggered from the Streamlit sidebar or from the graph flag. It probes authoritative URLs such as USPSTF, ADA, ACC/AHA themes, MedlinePlus, AHRQ, and PMC-related pages.

What is implemented: URL freshness signals and local superseded_by handling.

What is ideal and not claimed done: automatic PDF ingestion of every new society guideline release, continuous watchers with human approval gates, and a fully versioned corpus with clinical governance audit trails. Those stay in the green column.
""",
            "That same honesty appears in deployment — Figure 9 compares today’s laptop demo to a production-oriented target.",
        ),
        (
            13,
            "Figure 9 — Deployment Architecture (Today vs Target)",
            "~1–1.5 minutes",
            """
Yellow path — Milestone 2 demo runtime: developer or classroom laptop; Python venv; streamlit run app.py; optional uvicorn for FastAPI; FAISS and JSON indexes on local disk; optional OPENAI_API_KEY; public seed KB only — no PHI.

Green path — ideal future: HIPAA-eligible cloud on Azure or AWS; managed vector and hybrid search; Azure OpenAI, Bedrock, or Claude-class clinical LLM hosting; React clinician console with auth; CI evaluation and observability; guideline watchers with approval; optional EHR/FHIR only under proper BAAs.

I’m not claiming the green path is built. I’m showing I understand what production would require beyond a course prototype.
""",
            "Now the required instructor deliverable: the AI Tech Stack table with yellow and green highlights.",
        ),
        (
            14,
            "AI Tech Stack — Implemented vs Ideal",
            "~2–2.5 minutes",
            """
This slide follows the AI Tech Stack.docx template. Legend: yellow means implemented in Milestone 2; green means ideal future architecture.

I’ll answer the three instructor questions directly.

Question one — what stack did I implement? Look at the yellow cells. Data and vectors: local JSON medical KB and FAISS. Coding and agents: VS Code/Cursor, terminal scripts, LangGraph with LangChain ecosystem. GenAI: Ollama-compatible local path and OpenAI-compatible API when keyed. Frontend: Streamlit. Hosting: localhost Streamlit and optional Uvicorn. Bonus yellow rows: TF-IDF plus FAISS retrieval, LangGraph CDSS graph, and local eval JSON for latency and citations.

Question two — what is the ideal stack? Green cells: managed clinical KB or warehouse options like Supabase, BigQuery, Azure AI Search, OpenSearch; Vertex, SageMaker, or Azure ML for MLOps; GitHub Copilot as an assist; Gemini, Bedrock, Azure OpenAI, Anthropic Claude as stronger hosted LLMs; React/Next and cloud hosting; neural embeddings, hybrid search, and re-rankers; hosted workflow orchestrators; LangSmith-style observability and continuous groundedness with clinician override logs.

Question three — why those ideal components? That’s the next slide: cost, complexity, scalability, maintainability, latency, and safety.

While you’re looking at the table, notice the bonus rows. The instructor invited additions. Embeddings/retrieval, CDSS orchestration, and eval/observability are where this project’s real engineering work sits — so I made them explicit.
""",
            "Here’s why green is not yellow yet — the trade-off matrix.",
        ),
        (
            15,
            "Why This Ideal Stack? — Engineering Trade-offs",
            "~2 minutes",
            """
This answers instructor question three using AI-engineering decision dimensions.

Cost: implemented local TF-IDF and optional LLM keeps the demo near zero dollars. Ideal paid embeddings, stronger LLMs, and cloud raise quality but add recurring cost and, for PHI, BAA complexity.

Complexity: one Python repo plus Streamlit is easy to run and grade. Ideal React, managed search, and EHR FHIR expand the integration surface dramatically.

Scalability: FAISS on disk is fine for a nineteen-document prototype. Multi-clinic load needs managed vector search and autoscaling APIs.

Maintainability: clear LangGraph nodes and seed JSON are inspectable for a class project. Ideal CI evaluation and guideline watchers reduce drift but need ops ownership.

Latency: local retrieval is fast after index warm-up; LLM is optional. Ideal re-rankers and always-on LLMs add latency unless you add caching and async jobs.

Safety and trust: citations, confidence, review flags, and a loud disclaimer are implemented. Ideal adds clinical validation, audit logging, and a red-team evaluation harness.

The story is not “local is forever best.” The story is “local is the right Milestone 2 choice under cost, reliability, and demo constraints — and I know what upgrades buy me later.”
""",
            "Given those choices, here are the data, model, and insights from actually running the system.",
        ),
        (
            16,
            "Data, Model, and Key Insights",
            "~1.5 minutes",
            """
Data: public-style seed KB in JSON — guideline summaries, MedlinePlus, AHRQ, PubMed/PMC-style notes, MedQuAD sample. Pipeline is load, clean, chunk, TF-IDF, FAISS. About 19 documents and 20 chunks indexed. Knowledge update registry exists. I do not claim bulk PMC crawl or EHR/FHIR PHI integration.

Model and agents: LangGraph with eight nodes covering patient through transparency. Optional OpenAI-compatible chat; offline extractive fallback. Streamlit UI plus optional FastAPI.

Key insights from building it: guideline-first ranking beats pure semantic retrieval for clinical trust. Superseded documents must be penalized or answers drift. On the demo evaluation set, citation precision reached 1.0 when answers cite retrieved [n] sources. Offline path is required for reliable classroom demos. And CDSS structure is more clinically usable than a plain chatbot Q&A box.
""",
            "Those insights connect directly to the evaluation metrics on the next slide.",
        ),
        (
            17,
            "Evaluation — Technical Proof Points",
            "~1.5–2 minutes",
            """
Evaluation is produced by scripts/evaluate.py into outputs/eval_results.json. This maps proposal metrics onto what the prototype can measure now.

Citation precision on the demo eval set: 1.00 — valid [n] citations over total citations.

Chunks per query: 5 after priority re-rank top-k.

Indexed corpus: 19 documents / 20 chunks.

Retrieval path: FAISS with TF-IDF and hierarchy re-rank.

Qualitative RAG versus plain LLM comparison matches the proposal plan. With RAG: source links, groundedness constrained to retrieved context, agent-trace audit trail, and offline demo capability. Without retrieval, a fluent model can invent guidelines and has no retrieval audit — and without an API key, plain LLM isn’t even available. That contrast is exactly why RAG belongs in a clinical evidence assistant.
""",
            "Building this wasn’t frictionless — challenges and learnings next.",
        ),
        (
            18,
            "Challenges and Learnings",
            "~1–1.5 minutes",
            """
Challenges: heavy embedding stacks broke local installs, so I defaulted to offline TF-IDF. Pure semantic search overweighted patient-education pages. Superseded guidelines still retrieved unless ranking penalized them. The UI had to become clinician-readable without rewriting agents. I had to label indexed or cached evidence honestly versus true live retrieval. And medical prototypes need loud disclaimers and review flags.

Learnings: separate implemented versus ideal stack explicitly. Evidence hierarchy plus citation checks beat fluency-only answers. LangGraph boundaries make CDSS steps demoable and debuggable. Local-first design improves reliability and lowers cost. Trade-offs are the product story — cost, latency, scale, maintainability. And clinician support is not clinician replacement; transparency builds trust.
""",
            "Now the proof — live demo with camera on, walking the running system and the codebase.",
        ),
        (
            19,
            "Live Demo — Proof of Implementation",
            "~3–5 minutes (demo time)",
            """
For the recording: camera on.

Demo path: from AI HEALTH CARE, build the index if needed, then streamlit run app.py — or START_DEMO.bat.

I’ll enter patient fields and a clinical question — for example ASCVD statin intensity or A1C target.

I’ll run the CDSS workflow and show KPIs, assessment, risk, evidence tiles with guidelines first, verification, recommendations, citations, agent trace, and confidence.

I’ll point to Knowledge Update and Rebuild index as maintenance actions.

I’ll briefly show the codebase — graph.py, evidence_rank.py, vectorstore.py — so it’s clear this isn’t slides-only.

Honesty check while demoing: implemented are Streamlit, LangGraph CDSS graph, FAISS, TF-IDF, priority ranking, citations, optional FastAPI, and knowledge-update probes. Not done: managed vector DB, neural embeddings plus re-ranker, HIPAA cloud, React app, full PMC or EHR.

And I’ll repeat: research prototype, not a medical device; the clinician decides.
""",
            "I’ll close with the core takeaway.",
        ),
        (
            20,
            "Closing",
            "~30–45 seconds",
            """
Better clinical decisions need evidence clinicians can verify.

This Milestone 2 system demonstrates AI engineering for that goal: RAG, agents, ranking, and transparency — implemented, measured, and demoable.

I’m Audrey Rah, University of Houston. Happy to take questions.
""",
            "",
        ),
    ]

    for item in slides:
        add_slide(doc, *item)

    # Appendix: Q&A prep
    add_heading_styled(doc, "Appendix — Likely questions and concise answers", level=1)

    qa = [
        (
            "Why TF-IDF instead of neural embeddings?",
            "Reliability and cost for Milestone 2. Heavy embedding stacks broke installs. TF-IDF keeps the demo deterministic and offline. Ideal path adds neural embeddings and a re-ranker when environment and budget allow.",
        ),
        (
            "Does the system replace doctors?",
            "No. It is advisory CDSS support with citations, confidence, and human-review flags. Final decisions stay with a licensed clinician. Prototype is not a medical device.",
        ),
        (
            "Is live PubMed / full guideline PDF ingestion implemented?",
            "There is an optional PubMed helper and a Knowledge Update URL-probe prototype. Bulk PMC crawl and automatic society PDF ingestion are ideal/future — not claimed as complete.",
        ),
        (
            "How do you prevent hallucinations?",
            "Retrieve first, constrain answers to retrieved context, check citation indices, flag low confidence or conflicts for human review, and prefer guideline-tier evidence.",
        ),
        (
            "Where is the code?",
            "AI HEALTH CARE/ under the Final project folder — app.py, src/graph.py, evidence_rank.py, vectorstore.py, and related modules.",
        ),
    ]
    for q, a in qa:
        add_para(doc, f"Q: {q}", size=11, bold=True, space_after=2)
        add_para(doc, f"A: {a}", size=11, space_after=10)

    add_heading_styled(doc, "Recording checklist", level=1)
    for line in [
        "Camera ON for all presenters.",
        "Open Milestone2_HealthcareAI_V2.pptx (or the original Milestone 2 deck).",
        "Have AI HEALTH CARE Streamlit running before you start the demo slide.",
        "State disclaimer early and again during demo.",
        "When showing the tech stack, explicitly say yellow vs green.",
        "Do not claim green-cell components as finished.",
        "End with questions.",
    ]:
        add_para(doc, f"• {line}", size=11, space_after=4)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
